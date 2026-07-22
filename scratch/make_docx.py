import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    # Primary: Cyan/Navy (#0F172A, #0284C7, #38BDF8)
    COLOR_PRIMARY = RGBColor(15, 23, 42)     # Dark Slate
    COLOR_ACCENT = RGBColor(2, 132, 199)    # Deep Cyan
    COLOR_TEXT = RGBColor(51, 65, 85)       # Slate text
    COLOR_MUTED = RGBColor(100, 116, 139)   # Muted gray

    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_icon = p_title.add_run("🛡️\n")
    run_icon.font.size = Pt(36)
    
    run_title = p_title.add_run("ScamShield AI")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Digital Public Safety Intelligence Platform\nTechnical & Executive Specification")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT
        return h

    def add_bullet(p_or_text, level=0):
        if isinstance(p_or_text, str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(p_or_text)
            return p
        return p_or_text

    def set_cell_background(cell, fill_color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color_hex}"/>')
        tcPr.append(shd)

    # 1. Executive Summary
    add_heading_1("1. Executive Summary & Problem Analysis")
    doc.add_paragraph(
        "Digital fraud in India has evolved into an organized criminal enterprise. The Indian Cyber Crime "
        "Coordination Centre (I4C) reported over 694,000 cybercrime complaints in 2024, representing "
        "financial losses exceeding ₹11,333 Crore ($1.35 Billion USD) in the first six months alone."
    )
    doc.add_paragraph(
        "Among all fraud vectors, the 'Digital Arrest' scam is the most insidious. Perpetrators impersonate "
        "officials from the Central Bureau of Investigation (CBI), Narcotics Control Bureau (NCB), Enforcement "
        "Directorate (ED), or State Police departments. Victims are coerced into video calls, told that their identity "
        "or bank accounts are linked to illegal activities, and forced into continuous monitoring under threat of "
        "immediate physical arrest."
    )

    # Quantitative Fraud Impact Table
    add_heading_2("Quantitative Fraud Impact Matrix")
    table1 = doc.add_table(rows=5, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Fraud Vector", "Annual Case Count", "Avg Loss per Victim", "Psychological Mechanism"]
    for i, h in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F172A")

    data1 = [
        ["Digital Arrest", "7,061 (2024)", "₹15.5 Lakhs", "Coercive fear & authority pressure"],
        ["UPI / Payment Trap", "312,000 (2024)", "₹42,000", "Speed urgency & cognitive overload"],
        ["Phishing / Impersonation", "185,000 (2024)", "₹88,000", "Trust exploitation & deceptive URLs"],
        ["Investment / Stock Scams", "98,000 (2024)", "₹4.2 Lakhs", "Greed amplification & fake returns"]
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table1.cell(row_idx, col_idx)
            cell.text = text
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            else:
                set_cell_background(cell, "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. Biological Context
    add_heading_1("2. Biological & Societal Context of Digital Arrest Scams")
    doc.add_paragraph(
        "The success of digital arrest scams relies on exploiting human stress neurobiology. "
        "When a scammer claims to be a police officer threatening immediate imprisonment:"
    )
    add_bullet("1. The victim's amygdala triggers a sympathetic nervous system fight-or-flight response.")
    add_bullet("2. Adrenaline and cortisol surge, suppressing executive functioning in the prefrontal cortex.")
    add_bullet("3. Rational scrutiny drops significantly, making the victim highly susceptible to instructions such as transferring funds to 'government verification accounts'.")
    doc.add_paragraph(
        "ScamShield AI acts as an artificial prefrontal cortex guardrail. By displaying objective, "
        "mathematical risk meters and visual alerts, the platform interrupts the panic loop and forces analytical cognition to re-engage."
    )

    # 3. System Architecture
    add_heading_1("3. System Architecture & Data Engineering")
    doc.add_paragraph(
        "ScamShield AI is engineered as a decoupled, multi-tiered architecture designed for low latency, "
        "high throughput, and zero single-point-of-failure operation."
    )
    add_bullet("Frontend Layer: Built with Next.js 14 (App Router), TypeScript, and Tailwind CSS with custom glassmorphism styles.")
    add_bullet("Backend Layer: Built with Python 3 and FastAPI, featuring asynchronous processing, rate limiting, and security middleware.")
    add_bullet("Real-Time Communication Layer: Bi-directional WebSocket channels for streaming speech transcripts and real-time risk scores.")
    add_bullet("AI Engine Layer: Multi-model pipeline featuring Google Gemini (primary), Hugging Face (secondary), and a local deterministic heuristic engine (fallback).")

    # 4. AI Engine & Real-Time NLP Pipeline
    add_heading_1("4. AI Engine & Real-Time NLP Pipeline")
    doc.add_paragraph(
        "The ScamAnalyzer service implements a three-tier fallback chain. If cloud APIs are unreachable, "
        "the local heuristic engine scores inputs based on four domain risk vectors:"
    )
    add_bullet("Authority Impersonation Vector (Weight 30%): Matches phrases like CBI, Police, Customs, Supreme Court.")
    add_bullet("Coercion & Legal Threats Vector (Weight 35%): Matches phrases like digital arrest, non-bailable warrant, Section 420.")
    add_bullet("Financial Demand Vector (Weight 25%): Matches phrases like verification account, UPI deposit, security fee.")
    add_bullet("Urgency & Secrecy Vector (Weight 10%): Matches phrases like immediate, keep quiet, do not disconnect.")

    # 5. Graph Theory Mathematics
    add_heading_1("5. Graph Theory & Threat Intelligence Mathematics")
    doc.add_paragraph(
        "The Threat Intelligence Dashboard incorporates an interactive force-directed graph to map fraud networks. "
        "Node coordinates are dynamically calculated using Coulomb repulsion and Hooke spring attraction physics:"
    )
    doc.add_paragraph(
        "Repulsive Force (Coulomb's Law): Fr(i, j) = - (kr / ||r_ij||^2) * r_hat_ij\n"
        "Attractive Force (Hooke's Law): Fa(i, j) = ka * (||r_ij|| - l0) * r_hat_ij"
    )
    doc.add_paragraph(
        "Where kr = 5000 is the repulsion constant, ka = 0.05 is spring stiffness, and l0 = 100 is edge rest length."
    )

    # 6. Gamification Engine
    add_heading_1("6. Gamification & Behavior Change Engine")
    doc.add_paragraph(
        "To transform passive users into active fraud detectors, ScamShield AI implements an incentive mechanics loop:"
    )
    add_bullet("Scan Message: +10 Points")
    add_bullet("Share Fraud Alert on WhatsApp: +15 Points")
    add_bullet("Quiz Correct Answer: +20 Points")
    add_bullet("Complete Full Quiz: +50 Points")
    add_bullet("Submit Verified Report to Authorities: +25 Points")

    # 7. Security Hardening
    add_heading_1("7. Security Hardening & Threat Model")
    doc.add_paragraph(
        "ScamShield AI enforces strict HTTP security headers (X-Frame-Options DENY, X-Content-Type-Options nosniff, "
        "Permissions-Policy, HSTS), origin-restricted CORS policies, per-IP rate limiting (SlowAPI), "
        "and short-lived JWT tokens for WebSocket connections."
    )

    # Save document
    out_path = r"C:\Users\KIIT0001\Suryansh's Desktop\Projects\ScamShield\ScamShield_AI_Hackathon_Specification.docx"
    doc.save(out_path)
    print(f"Document created successfully at: {out_path}")

if __name__ == "__main__":
    create_document()
